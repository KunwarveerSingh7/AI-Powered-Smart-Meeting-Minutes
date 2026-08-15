from io import BytesIO
import unittest

from docx import Document
from pypdf import PdfWriter

from backend.document_processor import (
    EmptyDocumentError,
    UnsupportedFileTypeError,
    extract_text_from_bytes,
)


class DocumentProcessorTests(unittest.TestCase):
    def test_extracts_utf8_txt(self):
        content = (
            b"Project Meeting\n"
            b"Decision: Approve the database\n"
            b"Task: Complete document processing"
        )

        result = extract_text_from_bytes("meeting.txt", content)

        self.assertIn("Project Meeting", result)
        self.assertIn("Decision: Approve the database", result)
        self.assertIn("Task: Complete document processing", result)

    def test_extracts_cp1252_txt(self):
        content = "Manager’s meeting notes".encode("cp1252")

        result = extract_text_from_bytes("meeting.txt", content)

        self.assertEqual(result, "Manager’s meeting notes")

    def test_extracts_docx_paragraphs_and_tables(self):
        stream = BytesIO()
        document = Document()
        document.add_paragraph("Weekly Project Meeting")

        table = document.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "Decision"
        table.cell(0, 1).text = "Approve database schema"

        document.save(stream)

        result = extract_text_from_bytes(
            "meeting.docx",
            stream.getvalue(),
        )

        self.assertIn("Weekly Project Meeting", result)
        self.assertIn(
            "Decision | Approve database schema",
            result,
        )

    def test_rejects_unsupported_file_type(self):
        with self.assertRaises(UnsupportedFileTypeError):
            extract_text_from_bytes(
                "meeting.xlsx",
                b"example content",
            )

    def test_rejects_empty_file(self):
        with self.assertRaises(EmptyDocumentError):
            extract_text_from_bytes("meeting.txt", b"")

    def test_rejects_pdf_without_readable_text(self):
        stream = BytesIO()
        writer = PdfWriter()
        writer.add_blank_page(width=612, height=792)
        writer.write(stream)

        with self.assertRaises(EmptyDocumentError):
            extract_text_from_bytes(
                "blank.pdf",
                stream.getvalue(),
            )


if __name__ == "__main__":
    unittest.main()

from io import BytesIO
from pathlib import Path
import sys
import tempfile
import unittest

from docx import Document
from fastapi import FastAPI
from fastapi.testclient import TestClient
from sqlalchemy import create_engine
from sqlalchemy.orm import sessionmaker
from sqlalchemy.pool import StaticPool


PROJECT_ROOT = Path(__file__).resolve().parent.parent
BACKEND_DIRECTORY = PROJECT_ROOT / "backend"

if str(BACKEND_DIRECTORY) not in sys.path:
    sys.path.insert(0, str(BACKEND_DIRECTORY))

import meeting_router
import models
from database import Base


class MeetingRouterTests(unittest.TestCase):
    def setUp(self):
        self.temporary_upload_directory = tempfile.TemporaryDirectory()
        meeting_router.UPLOAD_DIRECTORY = Path(
            self.temporary_upload_directory.name
        )

        self.engine = create_engine(
            "sqlite://",
            connect_args={"check_same_thread": False},
            poolclass=StaticPool,
        )

        self.SessionLocal = sessionmaker(
            autocommit=False,
            autoflush=False,
            bind=self.engine,
        )

        Base.metadata.create_all(bind=self.engine)

        db = self.SessionLocal()
        db.add_all(
            [
                models.User(
                    email="manager@example.com",
                    hashed_password="test-password-hash",
                    role="manager",
                ),
                models.User(
                    email="employee@example.com",
                    hashed_password="test-password-hash",
                    role="employee",
                ),
            ]
        )
        db.commit()
        db.close()

        self.current_user = {
            "email": "manager@example.com",
            "role": "manager",
        }

        app = FastAPI()
        app.include_router(meeting_router.router)

        def override_get_db():
            db = self.SessionLocal()
            try:
                yield db
            finally:
                db.close()

        def override_current_user():
            return self.current_user

        app.dependency_overrides[
            meeting_router.get_db
        ] = override_get_db

        app.dependency_overrides[
            meeting_router.get_current_user
        ] = override_current_user

        self.client = TestClient(app)

    def tearDown(self):
        self.client.close()
        Base.metadata.drop_all(bind=self.engine)
        self.engine.dispose()
        self.temporary_upload_directory.cleanup()

    def test_manager_can_upload_txt_meeting(self):
        response = self.client.post(
            "/meetings/upload",
            data={"title": "Weekly Project Meeting"},
            files={
                "file": (
                    "meeting-notes.txt",
                    b"Decision: Approve the database\nTask: Test uploads",
                    "text/plain",
                )
            },
        )

        self.assertEqual(response.status_code, 201)

        response_data = response.json()

        self.assertEqual(
            response_data["title"],
            "Weekly Project Meeting",
        )
        self.assertEqual(
            response_data["original_filename"],
            "meeting-notes.txt",
        )
        self.assertEqual(response_data["file_type"], "txt")
        self.assertEqual(response_data["status"], "draft")
        self.assertIn(
            "Decision: Approve the database",
            response_data["extracted_text"],
        )

        db = self.SessionLocal()
        meeting = db.query(models.Meeting).one()

        self.assertEqual(meeting.uploaded_by, 1)
        self.assertEqual(
            meeting.raw_text,
            response_data["extracted_text"],
        )

        stored_path = Path(meeting.stored_file_path)
        self.assertTrue(stored_path.exists())
        self.assertEqual(
            stored_path.parent,
            Path(self.temporary_upload_directory.name),
        )
        self.assertNotEqual(
            stored_path.name,
            "meeting-notes.txt",
        )

        db.close()

    def test_employee_cannot_upload_meeting(self):
        self.current_user = {
            "email": "employee@example.com",
            "role": "employee",
        }

        response = self.client.post(
            "/meetings/upload",
            data={"title": "Employee Meeting"},
            files={
                "file": (
                    "meeting.txt",
                    b"Meeting notes",
                    "text/plain",
                )
            },
        )

        self.assertEqual(response.status_code, 403)
        self.assertEqual(
            response.json()["detail"],
            "Only managers can upload meeting documents.",
        )

    def test_rejects_unsupported_file_type(self):
        response = self.client.post(
            "/meetings/upload",
            data={"title": "Invalid Upload"},
            files={
                "file": (
                    "meeting.xlsx",
                    b"Spreadsheet content",
                    "application/vnd.ms-excel",
                )
            },
        )

        self.assertEqual(response.status_code, 415)

    def test_rejects_empty_document(self):
        response = self.client.post(
            "/meetings/upload",
            data={"title": "Empty Meeting"},
            files={
                "file": (
                    "meeting.txt",
                    b"",
                    "text/plain",
                )
            },
        )

        self.assertEqual(response.status_code, 400)

    def test_rejects_blank_title(self):
        response = self.client.post(
            "/meetings/upload",
            data={"title": "   "},
            files={
                "file": (
                    "meeting.txt",
                    b"Meeting notes",
                    "text/plain",
                )
            },
        )

        self.assertEqual(response.status_code, 422)


    @staticmethod
    def _build_text_pdf() -> bytes:
        """Create a small machine-readable PDF without extra dependencies."""
        content = (
            b"BT\n"
            b"/F1 12 Tf\n"
            b"72 720 Td\n"
            b"(Project PDF Meeting) Tj\n"
            b"0 -20 Td\n"
            b"(Decision: Approve PDF upload testing) Tj\n"
            b"ET\n"
        )

        objects = [
            b"<< /Type /Catalog /Pages 2 0 R >>",
            b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
            (
                b"<< /Type /Page /Parent 2 0 R "
                b"/MediaBox [0 0 612 792] "
                b"/Resources << /Font << /F1 4 0 R >> >> "
                b"/Contents 5 0 R >>"
            ),
            (
                b"<< /Type /Font /Subtype /Type1 "
                b"/BaseFont /Helvetica >>"
            ),
            (
                b"<< /Length "
                + str(len(content)).encode("ascii")
                + b" >>\nstream\n"
                + content
                + b"endstream"
            ),
        ]

        pdf = bytearray(b"%PDF-1.4\n")
        offsets = [0]

        for number, obj in enumerate(objects, start=1):
            offsets.append(len(pdf))
            pdf.extend(
                f"{number} 0 obj\n".encode("ascii")
            )
            pdf.extend(obj)
            pdf.extend(b"\nendobj\n")

        xref_offset = len(pdf)

        pdf.extend(
            f"xref\n0 {len(objects) + 1}\n".encode("ascii")
        )
        pdf.extend(b"0000000000 65535 f \n")

        for offset in offsets[1:]:
            pdf.extend(
                f"{offset:010d} 00000 n \n".encode("ascii")
            )

        pdf.extend(
            (
                f"trailer\n"
                f"<< /Size {len(objects) + 1} /Root 1 0 R >>\n"
                f"startxref\n{xref_offset}\n"
                f"%%EOF\n"
            ).encode("ascii")
        )

        return bytes(pdf)

    def test_manager_can_upload_pdf_meeting(self):
        response = self.client.post(
            "/meetings/upload",
            data={"title": "PDF Project Meeting"},
            files={
                "file": (
                    "meeting.pdf",
                    self._build_text_pdf(),
                    "application/pdf",
                )
            },
        )

        self.assertEqual(response.status_code, 201)

        data = response.json()

        self.assertEqual(data["file_type"], "pdf")
        self.assertIn(
            "Project PDF Meeting",
            data["extracted_text"],
        )
        self.assertIn(
            "Decision: Approve PDF upload testing",
            data["extracted_text"],
        )


    def test_manager_can_upload_docx_meeting(self):
        stream = BytesIO()
        document = Document()
        document.add_paragraph("Project Planning Meeting")
        document.add_paragraph("Decision: Complete backend integration")
        document.save(stream)

        response = self.client.post(
            "/meetings/upload",
            data={"title": "Project Planning Meeting"},
            files={
                "file": (
                    "meeting.docx",
                    stream.getvalue(),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )

        self.assertEqual(response.status_code, 201)

        data = response.json()

        self.assertEqual(data["file_type"], "docx")
        self.assertIn(
            "Decision: Complete backend integration",
            data["extracted_text"],
        )

    def test_rejects_file_larger_than_10_mb(self):
        oversized_content = b"A" * (
            meeting_router.MAX_UPLOAD_BYTES + 1
        )

        response = self.client.post(
            "/meetings/upload",
            data={"title": "Oversized Meeting"},
            files={
                "file": (
                    "meeting.txt",
                    oversized_content,
                    "text/plain",
                )
            },
        )

        self.assertEqual(response.status_code, 413)
        self.assertEqual(
            response.json()["detail"],
            "The uploaded file must not exceed 10 MB.",
        )


if __name__ == "__main__":
    unittest.main()

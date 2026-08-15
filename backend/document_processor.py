from __future__ import annotations

from io import BytesIO
from pathlib import Path
from typing import Final

from docx import Document
from pypdf import PdfReader


SUPPORTED_EXTENSIONS: Final[set[str]] = {".pdf", ".docx", ".txt"}


class DocumentProcessingError(ValueError):
    """Base error raised when an uploaded document cannot be processed."""


class UnsupportedFileTypeError(DocumentProcessingError):
    """Raised when the uploaded file type is not supported."""


class EmptyDocumentError(DocumentProcessingError):
    """Raised when no readable text can be extracted from a document."""


def _normalise_text(text: str) -> str:
    """Remove unnecessary whitespace while preserving readable line breaks."""
    cleaned_lines = [
        " ".join(line.split())
        for line in text.splitlines()
        if line.strip()
    ]
    return "\n".join(cleaned_lines).strip()


def _extract_pdf(content: bytes) -> str:
    reader = PdfReader(BytesIO(content))
    pages: list[str] = []

    for page in reader.pages:
        page_text = page.extract_text() or ""
        if page_text.strip():
            pages.append(page_text)

    return "\n".join(pages)


def _extract_docx(content: bytes) -> str:
    document = Document(BytesIO(content))
    extracted_parts: list[str] = []

    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            extracted_parts.append(paragraph.text)

    # Include text stored inside DOCX tables.
    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip()
                for cell in row.cells
                if cell.text.strip()
            )
            if row_text:
                extracted_parts.append(row_text)

    return "\n".join(extracted_parts)


def _extract_txt(content: bytes) -> str:
    try:
        return content.decode("utf-8-sig")
    except UnicodeDecodeError:
        try:
            return content.decode("cp1252")
        except UnicodeDecodeError as exc:
            raise DocumentProcessingError(
                "The TXT file could not be decoded as readable text."
            ) from exc


def extract_text_from_bytes(filename: str, content: bytes) -> str:
    """
    Validate a document and extract readable text from PDF, DOCX or TXT bytes.
    """
    if not filename or not filename.strip():
        raise DocumentProcessingError("A filename is required.")

    extension = Path(filename).suffix.lower()

    if extension not in SUPPORTED_EXTENSIONS:
        supported = ", ".join(sorted(SUPPORTED_EXTENSIONS))
        raise UnsupportedFileTypeError(
            f"Unsupported file type '{extension or 'unknown'}'. "
            f"Supported file types are: {supported}."
        )

    if not content:
        raise EmptyDocumentError("The uploaded file is empty.")

    try:
        if extension == ".pdf":
            extracted_text = _extract_pdf(content)
        elif extension == ".docx":
            extracted_text = _extract_docx(content)
        else:
            extracted_text = _extract_txt(content)
    except DocumentProcessingError:
        raise
    except Exception as exc:
        raise DocumentProcessingError(
            f"The {extension} document could not be processed."
        ) from exc

    cleaned_text = _normalise_text(extracted_text)

    if not cleaned_text:
        raise EmptyDocumentError(
            "No readable text was found in the uploaded document. "
            "Scanned image-only documents are not supported."
        )

    return cleaned_text
